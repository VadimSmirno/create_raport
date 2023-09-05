
import random
import docx
from docx.shared import RGBColor
from docx.shared import Pt
from datetime import datetime
from logging_dir.log import logger
from datadase.request_database import get_data_user
from datadase.models import User
from itertools import cycle
import json
import codecs


def get_month_in_genitive(date_string):

    current_date = datetime.strptime(date_string, "%d-%m-%Y")

    month_cases = {
        "january": "января",
        "february": "февраля",
        "march": "марта",
        "april": "апреля",
        "may": "мая",
        "june": "июня",
        "july": "июля",
        "august": "августа",
        "september": "сентября",
        "october": "октября",
        "november": "ноября",
        "december": "декабря"
    }

    month_name = current_date.strftime("%B").lower()
    month_in_genitive = month_cases.get(month_name, "")
    return month_in_genitive


def format_itinerary(itinerary):
    result = ""
    for destination, transport in itinerary.items():
        city:str = destination
        result += f"{transport} транспортом от г. Вологда до  {city}"

    return result


def format_itinerary2(itinerary):
    result = ""
    previous_destination = ""
    for destination, transport in itinerary.items():
        city: str = destination
        if previous_destination:
            result += f", {transport} транспортом от  {previous_destination} до {city}"
        else:
            result += f"{transport} транспортом от г. Вологда до  {city}"
        previous_destination = city
    return result


def format_vacation_followers(dct):
    result = "отпуска следуют:"
    if 'wife' in dct:
        result += f" жена {dct['wife']};"
    if 'husband' in dct:
        result += f" муж {dct['husband']};"
    if 'son' in dct:
        result += f" сын {dct['son']};"
    if 'daughter' in dct:
        result += f" дочь {dct['daughter']};"
    result = result.rstrip(";")  # Удаление последнего символа ";"
    return result


def generation_raport(telegram_id):
    fonts = ["Denistina", "PF Scandal Pro Black", "Marianna", "Shlapak Script", 'Liana']
    random_font = random.choice(fonts)
    result: User = get_data_user(telegram_id)
    decoded_str = codecs.decode(result.raport_info_json, 'unicode_escape')
    data = json.loads(decoded_str)
    date_start_vacation = data['date_start_vacation']
    date2 = datetime.strptime(date_start_vacation, "%d-%m-%Y")
    length_of_service = datetime.now().year - result.service_start_date.year

    if data['departure'] == 'Без выезда':
        doc = docx.Document('document_generation/Рапорт 1ч без выезда без стажа.docx')
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(16)
        font.color.rgb = RGBColor(102 , 0, 255)  # (89, 131, 176)
        font.name = random_font  # 'Denistina 16' "PF Scandal Pro Black  -14" "Marianna 16" ''Shlapak Script''

        for paragraph in doc.paragraphs:
            if length_of_service < 10 and 'и дополнительный отпуск за стаж службы' in paragraph.text:
                paragraph.text = paragraph.text.replace('и дополнительный отпуск за стаж службы в   ФПС  в', '')
                if 'количестве  dey календарных дней' in paragraph.text:
                    paragraph.text = paragraph.text.replace('количестве  dey календарных дней', '')
            elif 15 > length_of_service >= 10 and 'dey' in paragraph.text:
                paragraph.text = paragraph.text.replace('dey', '5')
            elif 20 > length_of_service >= 15 and 'dey' in paragraph.text:
                paragraph.text = paragraph.text.replace('dey', '10')
            elif length_of_service >= 20 and 'dey' in paragraph.text:
                paragraph.text = paragraph.text.replace('dey', '15')

            if data['vacation_part'] == '1':
                date_finish_vacation = data['date_finish_vacation']
                date1 = datetime.strptime(date_finish_vacation, "%d-%m-%Y")
                count = date1 - date2
                count = str(count.days)
                if '30 календарных дней' in paragraph.text:
                    paragraph.text = paragraph.text.replace('30', count)
            elif data['vacation_part'] == '2':
                if 'первую часть' in paragraph.text:
                    paragraph.text = paragraph.text.replace('первую', 'вторую')
                if  'продолжительность 30' in paragraph.text:
                     paragraph.text = paragraph.text.replace('отпуска продолжительность 30 календарных дней', 'отпуска')
            else:
                if  'первую часть' in paragraph.text:
                    paragraph.text = paragraph.text.replace('первую часть', 'основной отпуск')
                if  'продолжительность 30' in paragraph.text:
                     paragraph.text = paragraph.text.replace('основного отпуска продолжительность 30 календарных дней', '')
            if 'за 2023 год , с 2 августа.2023 года' in paragraph.text or '2023' in paragraph.text:
                month = get_month_in_genitive(date_start_vacation)
                paragraph.text = paragraph.text.replace('2023', str(date2.year))
                paragraph.text = paragraph.text.replace('2 августа', f'{str(date2.day)} {month}')
            if data[
                'manager_name'] != "Балчугову В.В" and paragraph.text == "Начальнику главного управления":
                paragraph.text = paragraph.text.replace("Начальнику", 'ВрИО начальника')
            if data['manager_name'] != "Балчугову В.В" and 'генерал-майору' in paragraph.text:
                paragraph.text = paragraph.text.replace('генерал-майору', data['rung'])
            if data['manager_name'] != "Балчугову В.В" and 'Балчугову В.В.' in paragraph.text:
                paragraph.text = paragraph.text.replace('Балчугову В.В.', data['manager_name'])
            if '20.05.2023г' in paragraph.text:
                paragraph.text = paragraph.text.replace('20.05.2023', datetime.now().strftime("%d.%m.%Y"))

            if '8900 536 71 36' in paragraph.text:
                paragraph.text = paragraph.text.replace('8900 536 71 36', result.telephone_number)
            if 'Иванов Иван Иванович' in paragraph.text:
                paragraph.text = paragraph.text.replace('Иванов Иван Иванович',
                                                        f'{result.last_name} {result.first_name} {result.surname}')
            if 'Старший-инструктор ' in paragraph.text:
                paragraph.text = paragraph.text.replace('Старший-инструктор по вождению ', result.job_title)
            if '3 пожарно' in paragraph.text:
                paragraph.text = paragraph.text.replace('3', str(result.part_number))
            if 'старшина' in paragraph.text:
                paragraph.text = paragraph.text.replace('старшина', result.rank)
            if data['material_aid'] == 'нет' and 'Прошу выплатить мне' in paragraph.text:
                p = paragraph._element
                p.getparent().remove(p)
        doc.save(f'document_generation/{result.last_name} {result.first_name[0]}.{result.surname[0]}.docx')
    elif data['departure'] == 'С выездом':
        doc = docx.Document('document_generation/Рапорт 1ч с выездом без стажа.docx')
        style = doc.styles['Normal']
        font = style.font
        font.size = Pt(16)
        font.color.rgb = RGBColor(102, 0, 255)  # (89, 131, 176)
        font.name = random_font  # 'Denistina 16' "PF Scandal Pro Black  -14" "Marianna 16" ''Shlapak Script''
        for paragraph in doc.paragraphs:
            if length_of_service < 10 and 'и дополнительный отпуск за стаж службы' in paragraph.text:
                paragraph.text = paragraph.text.replace('и дополнительный отпуск за стаж службы в   ФПС  в', '')
                if 'количестве  dey календарных дней' in paragraph.text:
                    paragraph.text = paragraph.text.replace('количестве  dey календарных дней', '')
            elif 15 > length_of_service >= 10 and 'dey' in paragraph.text:
                paragraph.text = paragraph.text.replace('dey', '5')
            elif 20 > length_of_service >= 15 and 'dey' in paragraph.text:
                paragraph.text = paragraph.text.replace('dey', '10')
            elif length_of_service >= 20 and 'dey' in paragraph.text:
                paragraph.text = paragraph.text.replace('dey', '15')
            if data['vacation_part'] == '1':
                date_finish_vacation = data['date_finish_vacation']
                date1 = datetime.strptime(date_finish_vacation, "%d-%m-%Y")
                count = date1 - date2
                count = str(count.days)
                if '30 календарных дней' in paragraph.text:
                    paragraph.text = paragraph.text.replace('30', count)

            elif data['vacation_part'] == '2':
                if 'первую часть' in paragraph.text:
                    paragraph.text = paragraph.text.replace('первую', 'вторую')
                if  'продолжительность 30' in paragraph.text:
                     paragraph.text = paragraph.text.replace('отпуска продолжительность 30 календарных дней', 'отпуска')
            else:
                if  'первую часть' in paragraph.text:
                    paragraph.text = paragraph.text.replace('первую часть', 'основной отпуск')
                if  'продолжительность 30' in paragraph.text:
                     paragraph.text = paragraph.text.replace('основного отпуска продолжительность 30 календарных дней', '')
            if 'за 2023 год , с 2 августа.2023 года' in paragraph.text or '2023' in paragraph.text:
                month = get_month_in_genitive(date_start_vacation)
                paragraph.text = paragraph.text.replace('2023', str(date2.year))
                paragraph.text = paragraph.text.replace('2 августа', f'{str(date2.day)} {month}')
            if data[
                'manager_name'] != "Балчугову В.В" and paragraph.text == "Начальнику главного управления":
                paragraph.text = paragraph.text.replace("Начальнику", 'ВрИО начальника')
            if data['manager_name'] != "Балчугову В.В" and 'генерал-майору' in paragraph.text:
                paragraph.text = paragraph.text.replace('генерал-майору', data['rung'])
            if data['manager_name'] != "Балчугову В.В" and 'Балчугову В.В.' in paragraph.text:
                paragraph.text = paragraph.text.replace('Балчугову В.В.', data['manager_name'])
            if '8900 536 71 36' in paragraph.text:
                paragraph.text = paragraph.text.replace('8900 536 71 36', result.telephone_number)
            if '20.05.2023г' in paragraph.text:
                paragraph.text = paragraph.text.replace('20.05.2023', datetime.now().strftime("%d.%m.%Y"))
            if '3 пожарно' in paragraph.text:
                paragraph.text = paragraph.text.replace('3', str(result.part_number))
            if 'Старший-инструктор ' in paragraph.text:
                paragraph.text = paragraph.text.replace('Старший-инструктор по вождению ', result.job_title)
            if 'старшина' in paragraph.text:
                paragraph.text = paragraph.text.replace('старшина', result.rank)
            if 'Иванов Иван Иванович' in paragraph.text:
                paragraph.text = paragraph.text.replace('Иванов Иван Иванович',
                                                        f'{result.last_name} {result.first_name} {result.surname}')
            if 'буду проводить в городе Сочи' in paragraph.text:
                list_city = data['list_city']
                logger.info(list_city)
                list_city =iter(list_city.split(', '))
                region, city = next(list_city,''), next(list_city,'')
                paragraph.text = paragraph.text.replace('городе Сочи', f'{region} {city}')


            if 'железнодорожным транспортом' in paragraph.text:
                itinerary = data.get('itinerary')
                if len(itinerary) == 1:
                    itinerary = format_itinerary(itinerary)
                    paragraph.text = paragraph.text.replace('железнодорожным транспортом', itinerary)
                else:
                    itinerary = format_itinerary2(itinerary)
                    paragraph.text = paragraph.text.replace('железнодорожным транспортом', itinerary)
            if 'отпуска следуют:' in paragraph.text:
                family = data
                family_str = format_vacation_followers(family)
                if family_str.endswith('отпуска следуют:'):
                    p = paragraph._element
                    p.getparent().remove(p)
                else:
                    paragraph.text = paragraph.text.replace('отпуска следуют:', family_str)
            if data['material_aid'] == 'нет' and 'Прошу выплатить мне' in paragraph.text:
                p = paragraph._element
                p.getparent().remove(p)

        doc.save(f'document_generation/{result.last_name} {result.first_name[0]}.{result.surname[0]}.docx')

    return (f'{result.last_name} {result.first_name[0]}.{result.surname[0]}.docx', random_font)

# if __name__ == '__main__':
#     generation_raport(telegram_id='')
